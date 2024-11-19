import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
import os
from docx import Document
from docx.shared import Cm,Pt
import io
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION


doc = Document()

doc.add_heading('Resultados Rueda de la Vida', 0)

# Función para cada sección
def bienestar_social():
    
    questions = [
        "Contribuyo con tiempo y/o dinero a proyectos sociales y comunitarios.",
        "Estoy comprometido con una vida de voluntariado.",
        "Demuestro imparcialidad y justicia al tratar con personas.",
        "Tengo una red de amigos cercanos y/o familiares.",
        "Me interesan los demás, incluso aquellos de diferentes antecedentes que el mío.",
        "Puedo equilibrar mis propias necesidades con las necesidades de los demás.",
        "Puedo comunicarme y llevarme bien con una gran variedad de personas.",
        "Obedezco las leyes y reglas de nuestra sociedad.",
        "Soy una persona compasiva y trato de ayudar a los demás cuando puedo.",
        "Apoyo y ayudo con las reuniones sociales de la familia, el vecindario y el trabajo."
    ]
    values = []
    for question in questions:
        doc.add_paragraph(question)
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            doc.add_paragraph(value)
            value = 0
            values.append(value)
            
        elif value == "A veces":
            doc.add_paragraph(value)
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            doc.add_paragraph(value)
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

def bienestar_intelectual():
     
    questions = [
        "Estoy interesado en aprender cosas nuevas.",
        "Trato de mantenerme al tanto de los asuntos actuales a nivel local, nacional e internacional.",
        "Disfruto asistiendo a conferencias, obras de teatro, actuaciones musicales, museos, galerías y/o bibliotecas.",
        "Selecciono cuidadosamente películas y/o programas de televisión.",
        "Disfruto de actividades/juegos mentales creativos y estimulantes.",
        "Estoy contento con la cantidad y variedad que leo.",
        "Me esfuerzo por mejorar mis habilidades verbales y escritas.",
        "Un programa de educación continua es/será importante para mí en mi carrera.",
        "Puedo analizar, sintetizar y ver más de un lado de un problema.",
        "Me gusta participar en discusiones intelectuales."
    ]
    values = []
    for question in questions:
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            value = 0
            values.append(value)
        elif value == "A veces":
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

def bienestar_ocupacional():
    questions = [
        "Estoy contento con mi elección de carrera.",
        "Estoy buscando trabajar.",
        "Mis responsabilidades/obligaciones laborales/estudiantiles son consistentes con mis valores.",
        "Los pagos/ventajas en mi elección de campo profesional son consistentes con mis valores.",
        "Estoy contento con el equilibrio entre mi tiempo de trabajo y mi tiempo libre.",
        "Estoy contento con la cantidad de control que tengo en mi trabajo.",
        "Mi trabajo me da satisfacción y estimulación personal.",
        "Estoy contento con el crecimiento profesional que he tenido.",
        "Siento que mi trabajo me permite hacer una diferencia.",
        "Mi trabajo contribuye positivamente a mi bienestar."
    ]
    values = []
    for question in questions:
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            value = 0
            values.append(value)
        elif value == "A veces":
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

def bienestar_emocional():

    questions = [
        "Puedo desarrollar y mantener relaciones cercanas.",
        "Acepto la responsabilidad de mis acciones.",
        "Veo desafíos y cambios como oportunidades de crecimiento.",
        "Siento que tengo un control considerable sobre mí.",
        "Puedo reírme de la vida y de mí mismo.",
        "Me siento bien conmigo mismo.",
        "Soy capaz de lidiar adecuadamente con el estrés.",
        "Puedo reconocer mis deficiencias personales y trabajar en ellas.",
        "Tengo un fuerte sentido de optimismo de vida y uso mis pensamientos y actitudes en formas que afirmen la vida.",
        "Mi trabajo contribuye positivamente a mi bienestar general."
    ]
    values = []
    for question in questions:
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            value = 0
            values.append(value)
        elif value == "A veces":
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

def bienestar_espiritual():
    
    questions = [
        "Me siento cómodo con mi vida espiritual.",
        "Existe una relación directa entre mis valores y acciones.",
        "Cuando me deprimo o me frustro, mis creencias me ayudan.",
        "La oración, la meditación y/o reflexión son parte de mi vida.",
        "La vida es significativa para mí, y siento que tengo un propósito.",
        "Puedo hablar cómodamente de mis valores y creencias.",
        "Constantemente estoy tratando de crecer espiritualmente.",
        "Soy tolerante e intento aprender sobre las creencias de otros.",
        "Encuentro consuelo en la naturaleza y el entorno que me rodea.",
        "Practico gratitud regularmente en mi vida."
    ]
    values = []
    for question in questions:
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            value = 0
            values.append(value)
        elif value == "A veces":
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

def bienestar_financiero():
    
    questions = [
        "¿Tienes efectivo en tu bolsillo?",
        "¿Equilibras tu presupuesto?",
        "¿Conoces la cantidad total de la deuda que tienes?",
        "¿Administras bien tu tiempo?",
        "¿Tienes una cuenta propia de ahorro?",
        "¿Sabes cuánto hay en tu cuenta de ahorro?",
        "¿Tu situación financiera te genera estrés?",
        "¿Realizas alguna actividad que te genere ingresos?",
        "¿Tienes un plan financiero a largo plazo?",
        "¿Evalúas regularmente tu situación financiera?"
    ]
    values = []
    for question in questions:
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            value = 0
            values.append(value)
        elif value == "A veces":
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

def bienestar_fisico():
    
    questions = [
        "Hago ejercicio aeróbico (vigoroso, continuo) durante 20 o 30 min. al menos 3 veces por semana.",
        "Como fruta, vegetales y granos integrales todos los días.",
        "Evito los productos de tabaco.",
        "Uso un cinturón de seguridad mientras viajo en o conduzco un automóvil.",
        "Minimizo deliberadamente mi ingesta de colesterol, grasas y aceites.",
        "Evito tomar bebidas alcohólicas o no consumo más de una bebida por día.",
        "Tengo una cantidad adecuada de sueño.",
        "Tengo mecanismos adecuados para enfrentar el estrés.",
        "Mantengo un horario regular de inmunizaciones, exámenes físicos, chequeos dentales y autoexámenes.",
        "Mantengo un peso razonable, evitando los extremos de sobrepeso y bajo peso."
    ]
    values = []
    for question in questions:
        value = st.selectbox(question, options=["Selecciona una opción","Muy raramente", "A veces", "Casi siempre"])  # Selección de valores 0, 1 o 2
        if value == "Muy raramente":
            value = 0
            values.append(value)
        elif value == "A veces":
            value = 1
            values.append(value)
        elif value == "Casi siempre":
            value = 2
            values.append(value)
        else:
            values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nanmean([v for v in values if v is not None])

# Configuración de la aplicación Streamlit
def radar_chart(values):
    categories = ['Social', 'Intelectual', 'Ocupacional', 'Emocional', 'Espiritual', 'Financiero', 'Físico']
    N = len(categories)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()

    # Asegúrate de que los valores tengan la misma longitud
    values = values + [values[0]]  # Cierra el gráfico
    angles += angles[:1]  # Cierra el gráfico en el círculo

    # Comprobación de longitud
    if len(values) != len(angles):
        st.error("Error: Las dimensiones de los valores y ángulos no coinciden.")
        return None

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='skyblue', alpha=0.4)
    ax.plot(angles, values, color='blue', linewidth=2, linestyle='solid')
    ax.set_yticks(np.arange(0, 3, 1))
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=12)

    return fig

def feedback(scores):
    feedbacks = []
    for score in scores:
        if score <= 0.5:
            feedbacks.append("Bajo: Considera estrategias para mejorar tu bienestar en esta área.")
        elif score <= 1.5:
            feedbacks.append("Medio: Estás en un buen camino, pero hay áreas que podrías fortalecer.")
        else:
            feedbacks.append("Alto: ¡Excelente! Estás haciendo un gran trabajo en esta área.")
    return feedbacks

def resultado(results):
    resultados = []
    for score in results:
        resultados.append(score)
    return resultados

def imagen_gráfico(values):
    categories = ['Social', 'Intelectual', 'Ocupacional', 'Emocional', 'Espiritual', 'Financiero', 'Físico']
    N = len(categories)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()

    # Asegúrate de que los valores tengan la misma longitud
    values = values + [values[0]]  # Cierra el gráfico
    angles += angles[:1]  # Cierra el gráfico en el círculo

    # Comprobación de longitud
    if len(values) != len(angles):
        st.error("Error: Las dimensiones de los valores y ángulos no coinciden.")
        return None

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='skyblue', alpha=0.4)
    ax.plot(angles, values, color='blue', linewidth=2, linestyle='solid')
    ax.set_yticks(np.arange(0, 3, 1))
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=12)

    plt.savefig('grafica_resultados.png', format='png', bbox_inches='tight')
    plt.close()  # Cerrar la figura para liberar memoria
    return 'grafica_resultados.png'  # Devolver la ruta del archivo

def main():
    st.title('Rueda de la Vida')

    # Inicializa la sesión para el estado actual
    if 'current_page' not in st.session_state:
        st.session_state.current_page = 0  # Comienza en la página 0 (Bienvenida)

    # Inicializa las puntuaciones si no están definidas
    if 'scores' not in st.session_state:
        st.session_state.scores = [0] * 7  # Inicializa las puntuaciones a cero para 7 áreas

    # Define las páginas
    pages = [
        "Bienvenida",
        "Bienestar Social",
        "Bienestar Intelectual",
        "Bienestar Ocupacional",
        "Bienestar Emocional",
        "Bienestar Espiritual",
        "Bienestar Financiero",
        "Bienestar Físico",
        "Resultados"
    ]

 # Calcular cuántas secciones se han completado
    completed_sections = sum(1 for score in st.session_state.scores if score != 0)

    # Muestra el contenido de la página actual
    page = pages[st.session_state.current_page]

    if page == "Bienvenida":
        st.write("""
        Esta encuesta tiene como objetivo evaluar diferentes dimensiones de tu bienestar.
        A través de una serie de preguntas, podrás reflexionar sobre aspectos importantes de tu vida en las siguientes áreas:
        - Bienestar Social
        - Bienestar Intelectual
        - Bienestar Ocupacional
        - Bienestar Emocional
        - Bienestar Espiritual
        - Bienestar Financiero
        - Bienestar Físico

        Por favor, responde cada pregunta con sinceridad. Al final de la encuesta, podrás ver un gráfico que refleja tu bienestar en cada una de estas dimensiones.
        """)

    elif page == "Bienestar Social":
        st.header('Bienestar Social')
        st.session_state.scores[0] = bienestar_social()

    elif page == "Bienestar Intelectual":
        st.header('Bienestar Intelectual')
        st.session_state.scores[1] = bienestar_intelectual()

    elif page == "Bienestar Ocupacional":
        st.header('Bienestar Ocupacional')
        st.session_state.scores[2] = bienestar_ocupacional()

    elif page == "Bienestar Emocional":
        st.header('Bienestar Emocional')
        st.session_state.scores[3] = bienestar_emocional()

    elif page == "Bienestar Espiritual":
        st.header('Bienestar Espiritual')
        st.session_state.scores[4] = bienestar_espiritual()

    elif page == "Bienestar Financiero":
        st.header('Bienestar Financiero')
        st.session_state.scores[5] = bienestar_financiero()

    elif page == "Bienestar Físico":
        st.header('Bienestar Físico')
        st.session_state.scores[6] = bienestar_fisico()

    elif page == "Resultados":
        st.header('Resultados')
        if all(score > 0 for score in st.session_state.scores):
            st.pyplot(radar_chart(st.session_state.scores))

            st.write(f"Sección actual: {pages[st.session_state.current_page]}")
            st.write(f"Secciones completadas: {completed_sections} de {len(pages) - 2}")  # Excluye 'Resultados'

            st.link_button("MiTec","https://mitec.itesm.mx/?_gl=1*10jvpnu*_gcl_aw*R0NMLjE3Mjk1Mjg3MTEuQ2owS0NRanc5OWU0QmhEaUFSSXNBSVNFN1A4aUxFN1RPQndIdWRYbGp3bFpNS05rUjYwRVZkRmVBcHlzaEMtcWNOT3R3dG1kV3lRbDA5a2FBby1hRUFMd193Y0I.*_gcl_au*OTI3NDE3MTM2LjE3Mjk1Mjg2NjY.*_ga*MjExNTg0MjAyOS4xNjgwNzMyMDY4*_ga_D9LSDN87GD*MTczMDA1NzM4Ny40MC4wLjE3MzAwNTczODcuNjAuMC4zMzY4MDIzODI.")
            
            graph_path = imagen_gráfico(st.session_state.scores)  # Guarda la gráfica

            doc.add_heading("Resultados")
            doc.add_picture(graph_path)

            
            resultados_dados = resultado(st.session_state.scores)
            records = (
                (1, "Bienestar Social", resultados_dados[0]*10),
                (2, "Bienestar Intelectual", resultados_dados[1]*10),
                (3, "Bienestar Ocupacional", resultados_dados[2]*10),
                (4, "Bienestar Emocional", resultados_dados[3]*10),
                (5, "Bienestar Espiritual", resultados_dados[4]*10),
                (6, "Bienestar Financiero", resultados_dados[5]*10),
                (7, "Bienestar Físico", resultados_dados[6]*10)
            )

            doc.add_section(WD_SECTION.NEW_PAGE)

            # Luego, puedes agregar un título o texto en esa sección
            doc.add_paragraph("Tabla de Resultados")

            # Crear tabla en el documento
    
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = 'Table Grid'  # Puedes usar 'Table Grid', 'Light Shading', 'Colorful List', etc.

            # Encabezado de la tabla
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Index'
            hdr_cells[1].text = 'Tipo de Bienestar'
            hdr_cells[2].text = 'Resultado'

            # Formato del encabezado
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(12)

            # Colores de fondo
            hdr_cells[0]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7D7E8"/>'.format(nsdecls('w'))))
            hdr_cells[1]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7D7E8"/>'.format(nsdecls('w'))))
            hdr_cells[2]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7D7E8"/>'.format(nsdecls('w'))))

            # Rellenar la tabla con datos y aplicar formato
            for qty, id, desc in records:
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(qty)
                row_cells[1].text = str(id)
                row_cells[2].text = str(desc)

                # Alineación de las celdas y estilo de fuente
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.paragraphs[0].runs[0].font.size = Pt(11)

                # Alternar color de fondo para filas
                for i, cell in enumerate(row_cells):
                    if qty % 2 == 0:  # Cambia el color en filas pares
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E3F2FD"/>'.format(nsdecls('w'))))
                        hdr_cells = tabla.rows[0].cells
                        hdr_cells[0].text = 'Index'
                        hdr_cells[1].text = 'Tipo de Bienestar'
                        hdr_cells[2].text = 'Resultado'

            doc.add_heading("Retroalimentación")
            feedback_scores = feedback(st.session_state.scores)
            for category, score in zip(['Social', 'Intelectual', 'Ocupacional', 'Emocional', 'Espiritual', 'Financiero', 'Físico'], feedback_scores):
                doc.add_paragraph(f"{category}: {score}")

            doc_download = doc

            bio = io.BytesIO()
            doc_download.save(bio)
            if doc_download:
                st.download_button(
                    label="Descarga tus resultados Resultados",
                    data=bio.getvalue(),
                    file_name="Resultados.docx",
                    mime="docx"
                )
            
            if os.path.exists(graph_path):
                with open(graph_path, "rb") as file:
                    if st.download_button("Descargar imagen", file, "resultados.png", "image/png"):
                        st.success("Captura de pantalla guardada como 'resultados.png'.")
                    
            st.subheader("Recursos de Apoyo")
            st.link_button("TecExplorerMTY","https://tec.rs/TecExplorerMTY")
        else:
            st.warning("Por favor, completa todas las secciones antes de ver los resultados.")

    # Botones de navegación
    col1, col2 = st.columns(2)

    # Botón "Página Anterior"
    if st.session_state.current_page > 0:
        with col1:
            if st.button('Página Anterior'):
                st.session_state.current_page -= 1
                st.rerun()


    # Botón "Siguiente Página"
    if st.session_state.current_page < len(pages) - 1:
        with col2:
            if st.button('Siguiente Página'):
                st.session_state.current_page += 1
                st.rerun()
            
    if st.session_state.current_page < len(pages) - 1:
        st.write(f"Sección actual: {pages[st.session_state.current_page]}")
        st.write(f"Secciones completadas: {completed_sections} de {len(pages) - 2}")  # Excluye 'Resultados'


if __name__ == '__main__':
    main()
