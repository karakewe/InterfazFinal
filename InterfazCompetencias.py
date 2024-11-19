import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
import os
from docx import Document
from docx.shared import Cm,Pt
import io
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION

doc = Document()

doc.add_heading('Resultados Test de Competencias Transversales', 0)

# Función para cada sección
def autoconocimiento():
    
    questions = [
        "¿Con qué frecuencia reflexionas sobre tus fortalezas y áreas de mejora?",
        "¿Cuando enfrentas un desafío, ¿qué tan bien identificas tus reacciones emocionales?",
        "¿Qué tan bien puedes describir tus valores personales y cómo influyen en tus decisiones?"
    ]
    values = []

    
    for question in questions:     

        if question == questions[0]:
            opciones = ["Selecciona una opción","Nunca", "Raramente", "A veces", "Frecuentemente", "Siempre"]
            value = st.selectbox(question, options=["Selecciona una opción","Nunca", "Raramente", "A veces", "Frecuentemente", "Siempre"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "No soy consciente de mis reacciones emocionales",
                        "Me doy cuenta de mis emociones después del evento",
                        "A veces identifico mis emociones durante el evento",
                        "Suelo ser consciente de mis emociones mientras suceden",
                        "Siempre soy plenamente consciente de mis emociones en el momento"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "No soy consciente de mis reacciones emocionales",
                        "Me doy cuenta de mis emociones después del evento",
                        "A veces identifico mis emociones durante el evento",
                        "Suelo ser consciente de mis emociones mientras suceden",
                        "Siempre soy plenamente consciente de mis emociones en el momento"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                         "No puedo describir mis valores personales",
                        "Tengo una idea vaga de mis valores",
                        "Puedo describir algunos de mis valores y su influencia",
                        "Soy consciente de mis valores y cómo influyen en muchas de mis decisiones",
                        "Tengo claridad total sobre mis valores y cómo influyen en todas mis decisiones"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                         "No puedo describir mis valores personales",
                        "Tengo una idea vaga de mis valores",
                        "Puedo describir algunos de mis valores y su influencia",
                        "Soy consciente de mis valores y cómo influyen en muchas de mis decisiones",
                        "Tengo claridad total sobre mis valores y cómo influyen en todas mis decisiones"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

def inteligencia_social():
     
    questions = [
       "¿Qué tan bien te adaptas a diferentes entornos sociales y estilos de comunicación?",
       "¿Con qué frecuencia tomas en cuenta las emociones y perspectivas de otras personas en tus interacciones?",
       "¿Qué tan cómodo te sientes iniciando y manteniendo conversaciones con personas nuevas?"

    ]
    values = []
    for question in questions:
        if question == questions[0]:
            opciones = ["Selecciona una opción", 
                        "Nunca me adapto a diferentes entornos sociales",
                        "Me cuesta adaptarme a diferentes entornos sociales",
                        "A veces puedo adaptarme, pero con dificultad",
                        "Generalmente me adapto bien a diferentes entornos sociales",
                        "Siempre me adapto fácilmente a diferentes entornos sociales"]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                                                    "Nunca me adapto a diferentes entornos sociales",
                                                    "Me cuesta adaptarme a diferentes entornos sociales",
                                                    "A veces puedo adaptarme, pero con dificultad",
                                                    "Generalmente me adapto bien a diferentes entornos sociales",
                                                    "Siempre me adapto fácilmente a diferentes entornos sociales"]) 
            
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "Nunca considero las emociones y perspectivas de los demás",
                        "Raramente considero las emociones y perspectivas de los demás",
                        "A veces considero las emociones y perspectivas de los demás",
                        "Frecuentemente considero las emociones y perspectivas de los demás",
                        "Siempre considero las emociones y perspectivas de los demás"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Nunca considero las emociones y perspectivas de los demás",
                        "Raramente considero las emociones y perspectivas de los demás",
                        "A veces considero las emociones y perspectivas de los demás",
                        "Frecuentemente considero las emociones y perspectivas de los demás",
                        "Siempre considero las emociones y perspectivas de los demás"]) 
            
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                        "Muy incómodo, evito iniciar conversaciones",
                        "Incómodo, solo hablo cuando es necesario",
                        "Algo incómodo, pero lo hago ocasionalmente",
                        "Cómodo, inicio y mantengo conversaciones sin problemas",
                        "Muy cómodo, disfruto iniciar y mantener conversaciones con personas nuevas"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Muy incómodo, evito iniciar conversaciones",
                        "Incómodo, solo hablo cuando es necesario",
                        "Algo incómodo, pero lo hago ocasionalmente",
                        "Cómodo, inicio y mantengo conversaciones sin problemas",
                        "Muy cómodo, disfruto iniciar y mantener conversaciones con personas nuevas"]) 
            
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

def compromiso_etico():
    questions = [
       "¿Qué tan importante consideras la integridad académica (evitar el plagio, hacer trampa, etc.) en tus estudios?",
       "¿Cómo actúas cuando presencias una conducta que va en contra de tus principios éticos en un entorno académico o profesional?",
       "¿Con qué frecuencia reflexionas sobre las implicaciones éticas de tus decisiones y acciones?"
    ]
    values = []
    for question in questions:
        if question == questions[0]:
            opciones = ["Selecciona una opción",
                            "No es importante para mí",
                            "Es poco importante para mí",
                            "Es algo importante para mí",
                            "Es bastante importante para mí",
                            "Es extremadamente importante para mí"]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                                                        "No es importante para mí",
                                                        "Es poco importante para mí",
                                                        "Es algo importante para mí",
                                                        "Es bastante importante para mí",
                                                        "Es extremadamente importante para mí"]) 

            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "No hago nada, no es mi problema",
    "Raramente actúo, prefiero no involucrarme",
    "A veces actúo, pero depende de la situación",
    "Generalmente actúo, trato de abordar la situación",
    "Siempre actúo, defiendo mis principios éticos firmemente"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "No hago nada, no es mi problema",
    "Raramente actúo, prefiero no involucrarme",
    "A veces actúo, pero depende de la situación",
    "Generalmente actúo, trato de abordar la situación",
    "Siempre actúo, defiendo mis principios éticos firmemente"]) 
            
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                         "Nunca reflexiono sobre las implicaciones éticas",
    "Raramente reflexiono sobre las implicaciones éticas",
    "A veces reflexiono sobre las implicaciones éticas",
    "Frecuentemente reflexiono sobre las implicaciones éticas",
    "Siempre reflexiono sobre las implicaciones éticas antes de actuar"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                         "Nunca reflexiono sobre las implicaciones éticas",
    "Raramente reflexiono sobre las implicaciones éticas",
    "A veces reflexiono sobre las implicaciones éticas",
    "Frecuentemente reflexiono sobre las implicaciones éticas",
    "Siempre reflexiono sobre las implicaciones éticas antes de actuar"
                        ]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

def emprendimiento_innovador():

    questions = [
        "¿Con qué frecuencia propones nuevas ideas o soluciones para problemas en tu entorno académico o profesional?",
        "¿Qué tan dispuesto estás a asumir riesgos calculados para llevar a cabo un proyecto innovador?",
        "¿Cómo manejas los fracasos o contratiempos cuando trabajas en un proyecto innovador?"
        
    ]
    values = []
    for question in questions:
        if question == questions[0]:
            opciones = ["Selecciona una opción","Nunca propongo nuevas ideas",
    "Raramente propongo nuevas ideas",
    "A veces propongo nuevas ideas",
    "Frecuentemente propongo nuevas ideas",
    "Siempre estoy proponiendo nuevas ideas"
    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                                                    "Nunca propongo nuevas ideas",
    "Raramente propongo nuevas ideas",
    "A veces propongo nuevas ideas",
    "Frecuentemente propongo nuevas ideas",
    "Siempre estoy proponiendo nuevas ideas"
    ]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "Nunca estoy dispuesto a asumir riesgos",
    "Raramente estoy dispuesto a asumir riesgos",
    "A veces estoy dispuesto a asumir riesgos, pero con dudas",
    "Frecuentemente estoy dispuesto a asumir riesgos calculados",
    "Siempre estoy dispuesto a asumir riesgos calculados para innovar"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Nunca estoy dispuesto a asumir riesgos",
    "Raramente estoy dispuesto a asumir riesgos",
    "A veces estoy dispuesto a asumir riesgos, pero con dudas",
    "Frecuentemente estoy dispuesto a asumir riesgos calculados",
    "Siempre estoy dispuesto a asumir riesgos calculados para innovar"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                        "Me desanimo y dejo de intentar",
    "Raramente vuelvo a intentar después de un fracaso",
    "A veces me siento desmotivado, pero eventualmente vuelvo a intentar",
    "Generalmente aprendo del fracaso y busco nuevas formas de abordar el problema",
    "Siempre veo los fracasos como oportunidades de aprendizaje y persisto hasta lograr el objetivo"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Me desanimo y dejo de intentar",
    "Raramente vuelvo a intentar después de un fracaso",
    "A veces me siento desmotivado, pero eventualmente vuelvo a intentar",
    "Generalmente aprendo del fracaso y busco nuevas formas de abordar el problema",
    "Siempre veo los fracasos como oportunidades de aprendizaje y persisto hasta lograr el objetivo"
]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

def comunicacion():
    
    questions = [
        "¿Con qué claridad y efectividad expresas tus ideas en presentaciones orales o escritas?",
        "¿Cómo manejas las diferencias de opinión durante una discusión o trabajo en equipo?",
        "¿Con qué frecuencia solicitas y utilizas retroalimentación para mejorar tu forma de comunicarte?"    
    ]
    values = []
    for question in questions:
        if question == questions[0]:
            opciones = ["Selecciona una opción",
                        "Nunca logro expresar mis ideas claramente",
    "Raramente logro expresar mis ideas claramente",
    "A veces logro expresar mis ideas claramente",
    "Frecuentemente logro expresar mis ideas claramente",
    "Siempre logro expresar mis ideas claramente y efectivamente"]
            value = st.selectbox(question, options=["Selecciona una opción",
                                                    "Nunca logro expresar mis ideas claramente",
    "Raramente logro expresar mis ideas claramente",
    "A veces logro expresar mis ideas claramente",
    "Frecuentemente logro expresar mis ideas claramente",
    "Siempre logro expresar mis ideas claramente y efectivamente"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "Evito participar en discusiones con diferencias de opinión",
    "Raramente trato de entender otras perspectivas",
    "A veces intento entender otras perspectivas, pero me cuesta trabajo",
    "Generalmente escucho y considero las opiniones de los demás",
    "Siempre escucho activamente y manejo las diferencias de opinión de manera constructiva"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Evito participar en discusiones con diferencias de opinión",
    "Raramente trato de entender otras perspectivas",
    "A veces intento entender otras perspectivas, pero me cuesta trabajo",
    "Generalmente escucho y considero las opiniones de los demás",
    "Siempre escucho activamente y manejo las diferencias de opinión de manera constructiva"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                        "Nunca solicito retroalimentación",
    "Raramente solicito retroalimentación",
    "A veces solicito retroalimentación y la utilizo ocasionalmente",
    "Frecuentemente solicito retroalimentación y la utilizo para mejorar",
    "Siempre solicito retroalimentación y la incorporo para mejorar continuamente"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Nunca solicito retroalimentación",
    "Raramente solicito retroalimentación",
    "A veces solicito retroalimentación y la utilizo ocasionalmente",
    "Frecuentemente solicito retroalimentación y la utilizo para mejorar",
    "Siempre solicito retroalimentación y la incorporo para mejorar continuamente"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

def transformacion_digital():
    
    questions = [
        "¿Qué tan familiarizado estás con las herramientas digitales y tecnologías emergentes en tu campo de estudio?",
        "¿Cómo integras las herramientas digitales en tus trabajos académicos o proyectos?",
        "¿Qué tan dispuesto estás a aprender y adaptarte a nuevas tecnologías digitales que pueden mejorar tu rendimiento académico o profesional?"
    ]
    values = []
    for question in questions:
        if question == questions[0]:
            opciones = ["Selecciona una opción",
                         "No estoy familiarizado con ninguna herramienta digital o tecnología emergente",
    "Estoy poco familiarizado con algunas herramientas digitales",
    "Estoy algo familiarizado con algunas herramientas y tecnologías emergentes",
    "Estoy bastante familiarizado con varias herramientas y tecnologías emergentes",
    "Estoy muy familiarizado y utilizo regularmente herramientas digitales y tecnologías emergentes"]
            value = st.selectbox(question, options=["Selecciona una opción",
                                                     "No estoy familiarizado con ninguna herramienta digital o tecnología emergente",
    "Estoy poco familiarizado con algunas herramientas digitales",
    "Estoy algo familiarizado con algunas herramientas y tecnologías emergentes",
    "Estoy bastante familiarizado con varias herramientas y tecnologías emergentes",
    "Estoy muy familiarizado y utilizo regularmente herramientas digitales y tecnologías emergentes"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "Nunca utilizo herramientas digitales en mis trabajos o proyectos",
    "Raramente utilizo herramientas digitales, prefiero métodos tradicionales",
    "A veces utilizo herramientas digitales, pero de manera limitada",
    "Frecuentemente utilizo herramientas digitales para mejorar mis trabajos o proyectos",
    "Siempre integro herramientas digitales de manera eficiente y creativa en mis trabajos o proyectos"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Nunca utilizo herramientas digitales en mis trabajos o proyectos",
    "Raramente utilizo herramientas digitales, prefiero métodos tradicionales",
    "A veces utilizo herramientas digitales, pero de manera limitada",
    "Frecuentemente utilizo herramientas digitales para mejorar mis trabajos o proyectos",
    "Siempre integro herramientas digitales de manera eficiente y creativa en mis trabajos o proyectos"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                         "Nunca estoy dispuesto a aprender nuevas tecnologías",
    "Raramente estoy dispuesto a aprender nuevas tecnologías",
    "A veces estoy dispuesto a aprender, pero con cierta resistencia",
    "Generalmente estoy dispuesto a aprender y adaptarme a nuevas tecnologías",
    "Siempre estoy dispuesto y entusiasmado por aprender y adaptarme a nuevas tecnologías"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Nunca estoy dispuesto a aprender nuevas tecnologías",
    "Raramente estoy dispuesto a aprender nuevas tecnologías",
    "A veces estoy dispuesto a aprender, pero con cierta resistencia",
    "Generalmente estoy dispuesto a aprender y adaptarme a nuevas tecnologías",
    "Siempre estoy dispuesto y entusiasmado por aprender y adaptarme a nuevas tecnologías"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

def razonamiento_complejo():
    
    questions = [
        "¿Qué tan bien puedes identificar y analizar múltiples factores y perspectivas en situaciones complejas?",
        "¿Cuando enfrentas un problema complejo, ¿cómo manejas la incertidumbre y la ambigüedad?",
        "¿Qué tan efectivo eres en la elaboración de soluciones creativas y viables para problemas complejos?"
    ]
    values = []
    for question in questions:
        if question == questions[0]:
            opciones = ["Selecciona una opción",
                        "Nunca puedo identificar y analizar múltiples factores",
    "Raramente logro identificar y analizar múltiples factores",
    "A veces puedo identificar y analizar algunos factores y perspectivas",
    "Frecuentemente logro identificar y analizar múltiples factores y perspectivas",
    "Siempre puedo identificar y analizar múltiples factores y perspectivas en situaciones complejas"
    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                                                    "Nunca puedo identificar y analizar múltiples factores",
    "Raramente logro identificar y analizar múltiples factores",
    "A veces puedo identificar y analizar algunos factores y perspectivas",
    "Frecuentemente logro identificar y analizar múltiples factores y perspectivas",
    "Siempre puedo identificar y analizar múltiples factores y perspectivas en situaciones complejas"
    ]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras
        
        if question == questions[1]:
            opciones = ["Selecciona una opción",
                        "Me siento completamente abrumado y no sé cómo proceder",
    "Raramente puedo manejar la incertidumbre y la ambigüedad",
    "A veces puedo manejar la incertidumbre y la ambigüedad con dificultad",
    "Generalmente puedo manejar la incertidumbre y la ambigüedad de manera efectiva",
    "Siempre puedo manejar la incertidumbre y la ambigüedad con confianza y claridad"
                    ]
            
            value = st.selectbox(question, options=["Selecciona una opción",
                        "Me siento completamente abrumado y no sé cómo proceder",
    "Raramente puedo manejar la incertidumbre y la ambigüedad",
    "A veces puedo manejar la incertidumbre y la ambigüedad con dificultad",
    "Generalmente puedo manejar la incertidumbre y la ambigüedad de manera efectiva",
    "Siempre puedo manejar la incertidumbre y la ambigüedad con confianza y claridad"
    ]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

        if question == questions[2]:
            opciones = ["Selecciona una opción",
                         "Nunca puedo elaborar soluciones creativas y viables",
    "Raramente puedo elaborar soluciones creativas y viables",
    "A veces puedo elaborar soluciones creativas, pero son limitadas en viabilidad",
    "Frecuentemente puedo elaborar soluciones creativas y viables",
    "Siempre puedo elaborar soluciones creativas y viables para problemas complejos"
                    ]
            value = st.selectbox(question, options=["Selecciona una opción",
                         "Nunca puedo elaborar soluciones creativas y viables",
    "Raramente puedo elaborar soluciones creativas y viables",
    "A veces puedo elaborar soluciones creativas, pero son limitadas en viabilidad",
    "Frecuentemente puedo elaborar soluciones creativas y viables",
    "Siempre puedo elaborar soluciones creativas y viables para problemas complejos"]) 
            if value == opciones[1]:
                value = 1
                values.append(value)
                
            elif value == opciones[2]:
                value = 2
                values.append(value)

            elif value == opciones[3]:
                value = 3
                values.append(value)

            elif value == opciones[4]:
                value = 4
                values.append(value)

            elif value == opciones[5]:
                value = 5
                values.append(value)

            else:
                values.append(None)  # Manejar como prefieras

    # Retornar el promedio, omitiendo None
    return np.nansum([v for v in values if v is not None])

# Configuración de la aplicación Streamlit
def radar_chart(values):
    categories = ['Autoconocimiento', 'Inteligencia Social', 'Compromiso Ético', 'Emprendimiento Inn.', 'Comunicación', 'Transformación Dígital', 'Razonamiento Complejo']
    N = len(categories)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()

    # Asegúrate de que los valores tengan la misma longitud
    values = values + [values[0]]  # Cierra el gráfico
    angles += angles[:1]  # Cierra el gráfico en el círculo

    # Comprobación de longitud
    if len(values) != len(angles):
        st.error("Error: Las dimensiones de los valores y ángulos no coinciden.")
        return None

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='skyblue', alpha=0.4)
    ax.plot(angles, values, color='blue', linewidth=2, linestyle='solid')
    ax.set_yticks(np.arange(0, 16, 5))
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=8)
    ax.tick_params(axis='x', pad=28)

    return fig

def feedback(scores):
    feedbacks = []
    for score in scores:
        if score <= 5:
            feedbacks.append("Bajo: Considera estrategias para mejorar tu bienestar en esta área.")
        elif score <= 10:
            feedbacks.append("Medio: Estás en un buen camino, pero hay áreas que podrías fortalecer.")
        else:
            feedbacks.append("Alto: ¡Excelente! Estás haciendo un gran trabajo en esta área.")
    return feedbacks

def resultado(results):
    resultados = []
    for score in results:
        resultados.append(score)
    return resultados

def imagen_gráfico(values):
    categories = ['Autoconocimiento', 'Inteligencia Social', 'Compromiso Ético', 'Emprendimiento Inn.', 'Comunicación', 'Transformación Dígital', 'Razonamiento Complejo']
    N = len(categories)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()

    # Asegúrate de que los valores tengan la misma longitud
    values = values + [values[0]]  # Cierra el gráfico
    angles += angles[:1]  # Cierra el gráfico en el círculo

    # Comprobación de longitud
    if len(values) != len(angles):
        st.error("Error: Las dimensiones de los valores y ángulos no coinciden.")
        return None

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='skyblue', alpha=0.4)
    ax.plot(angles, values, color='blue', linewidth=2, linestyle='solid')
    ax.set_yticks(np.arange(0, 16, 1))
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=12)

    plt.savefig('grafica_resultados.png', format='png', bbox_inches='tight')
    plt.close()  # Cerrar la figura para liberar memoria
    return 'grafica_resultados.png'  # Devolver la ruta del archivo

def main():
    st.title('Test de Competencias Transversales')

    # Inicializa la sesión para el estado actual
    if 'current_page' not in st.session_state:
        st.session_state.current_page = 0  # Comienza en la página 0 (Bienvenida)

    # Inicializa las puntuaciones si no están definidas
    if 'scores' not in st.session_state:
        st.session_state.scores = [0] * 7  # Inicializa las puntuaciones a cero para 7 áreas

    # Define las páginas
    pages = [
        'Bienvenida',
        'Autoconocimiento', 
        'Inteligencia Social', 
        'Compromiso Ético', 
        'Emprendimiento Innovador', 
        'Comunicación', 
        'Transformación Dígital', 
        'Razonamiento Complejo',
        'Resultados'
    ]

 # Calcular cuántas secciones se han completado
    completed_sections = sum(1 for score in st.session_state.scores if score != 0)

    # Muestra el contenido de la página actual
    page = pages[st.session_state.current_page]

    if page == "Bienvenida":
        st.write("""
        Este test está diseñado para evaluar tus competencias en áreas clave de desarrollo personal y profesional. 
        A través de una serie de preguntas, podrás reflexionar sobre aspectos importantes de tus competencias en las siguientes áreas:
        - Autoconocimiento
        - Inteligencia Social
        - Compromiso Ético
        - Emprendimiento Innovador
        - Comunicación
        - Transformación Dígital
        - Razonamiento Complejo

        Por favor, responde cada pregunta con sinceridad. Al final de la encuesta, podrás ver un gráfico que refleja la evaluación de cada una de estas dimensiones.
        """)

    elif page == "Autoconocimiento":
        st.header('Autoconocimiento')
        st.session_state.scores[0] = autoconocimiento()

    elif page == "Inteligencia Social":
        st.header('Inteligencia Social')
        st.session_state.scores[1] = inteligencia_social()

    elif page == "Compromiso Ético":
        st.header('Compromiso Ético')
        st.session_state.scores[2] = compromiso_etico()

    elif page == "Emprendimiento Innovador":
        st.header('Emprendimiento Innovador')
        st.session_state.scores[3] = emprendimiento_innovador()

    elif page == "Comunicación":
        st.header('Comunicación')
        st.session_state.scores[4] = comunicacion()

    elif page == "Transformación Dígital":
        st.header('Transformación Dígital')
        st.session_state.scores[5] = transformacion_digital()

    elif page == "Razonamiento Complejo":
        st.header('Razonamiento Complejo')
        st.session_state.scores[6] = razonamiento_complejo()

    elif page == "Resultados":
        st.header('Resultados')
        if all(score > 0 for score in st.session_state.scores):
            st.pyplot(radar_chart(st.session_state.scores))

            st.write(f"Sección actual: {pages[st.session_state.current_page]}")
            st.write(f"Secciones completadas: {completed_sections} de {len(pages) - 2}")  
            
            st.link_button("MiTec","https://mitec.itesm.mx/?_gl=1*10jvpnu*_gcl_aw*R0NMLjE3Mjk1Mjg3MTEuQ2owS0NRanc5OWU0QmhEaUFSSXNBSVNFN1A4aUxFN1RPQndIdWRYbGp3bFpNS05rUjYwRVZkRmVBcHlzaEMtcWNOT3R3dG1kV3lRbDA5a2FBby1hRUFMd193Y0I.*_gcl_au*OTI3NDE3MTM2LjE3Mjk1Mjg2NjY.*_ga*MjExNTg0MjAyOS4xNjgwNzMyMDY4*_ga_D9LSDN87GD*MTczMDA1NzM4Ny40MC4wLjE3MzAwNTczODcuNjAuMC4zMzY4MDIzODI.")

            graph_path = imagen_gráfico(st.session_state.scores)  # Guarda la gráfica

            doc.add_heading("Resultados")
            doc.add_picture(graph_path)

            
            resultados_dados = resultado(st.session_state.scores)
            records = (
                (1, "Autoconocimiento", resultados_dados[0]),
                (2, "Inteligencia Social", resultados_dados[1]),
                (3, "Compromiso Ético", resultados_dados[2]),
                (4, "Emprendimiento Innovador", resultados_dados[3]),
                (5, "Comunicación", resultados_dados[4]),
                (6, "Transformación Dígital", resultados_dados[5]),
                (7, "Razonamiento Complejo", resultados_dados[6])
            )

            doc.add_section(WD_SECTION.NEW_PAGE)

            # Luego, puedes agregar un título o texto en esa sección
            doc.add_paragraph("Tabla de Resultados")

            # Crear tabla en el documento
    
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = 'Table Grid'  # Puedes usar 'Table Grid', 'Light Shading', 'Colorful List', etc.

            # Encabezado de la tabla
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Index'
            hdr_cells[1].text = 'Tipo de Competencia'
            hdr_cells[2].text = 'Resultado'

            # Formato del encabezado
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(12)

            # Colores de fondo
            hdr_cells[0]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7D7E8"/>'.format(nsdecls('w'))))
            hdr_cells[1]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7D7E8"/>'.format(nsdecls('w'))))
            hdr_cells[2]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7D7E8"/>'.format(nsdecls('w'))))

            # Rellenar la tabla con datos y aplicar formato
            for qty, id, desc in records:
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(qty)
                row_cells[1].text = str(id)
                row_cells[2].text = str(desc)

                # Alineación de las celdas y estilo de fuente
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.paragraphs[0].runs[0].font.size = Pt(11)

                # Alternar color de fondo para filas
                for i, cell in enumerate(row_cells):
                    if qty % 2 == 0:  # Cambia el color en filas pares
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E3F2FD"/>'.format(nsdecls('w'))))
                        hdr_cells = tabla.rows[0].cells
                        hdr_cells[0].text = 'Index'
                        hdr_cells[1].text = 'Tipo de Competencia'
                        hdr_cells[2].text = 'Resultado'

            doc.add_heading("Retroalimentación")
            feedback_scores = feedback(st.session_state.scores)
            for category, score in zip(['Autoconocimiento', 'Inteligencia Social', 'Compromiso Ético', 'Emprendimiento Innovador', 'Comunicación', 'Transformación Dígital', 'Razonamiento Complejo'], feedback_scores):
                doc.add_paragraph(f"{category}: {score}")

            doc_download = doc

            bio = io.BytesIO()
            doc_download.save(bio)
            if doc_download:
                st.download_button(
                    label="Descarga tus resultados Resultados",
                    data=bio.getvalue(),
                    file_name="Resultados.docx",
                    mime="docx"
                )
            
            if os.path.exists(graph_path):
                with open(graph_path, "rb") as file:
                    if st.download_button("Descargar imagen", file, "resultados.png", "image/png"):
                        st.success("Captura de pantalla guardada como 'resultados.png'.")

            st.subheader("Recursos de Apoyo")
            st.link_button("TecExplorerMTY","https://tec.rs/TecExplorerMTY")

        else:
            st.warning("Por favor, completa todas las secciones antes de ver los resultados.")

    # Botones de navegación
    col1, col2 = st.columns(2)

    # Botón "Página Anterior"
    if st.session_state.current_page > 0:
        with col1:
            if st.button('Página Anterior'):
                st.session_state.current_page -= 1
                st.rerun()


    # Botón "Siguiente Página"
    if st.session_state.current_page < len(pages) - 1:
        with col2:
            if st.button('Siguiente Página'):
                st.session_state.current_page += 1
                st.rerun()
            
    if st.session_state.current_page < len(pages) - 1:
        st.write(f"Sección actual: {pages[st.session_state.current_page]}")
        st.write(f"Secciones completadas: {completed_sections} de {len(pages) - 2}")  

if __name__ == '__main__':
    main()
